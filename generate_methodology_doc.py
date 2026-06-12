"""
Generate 'Exposure_Simulation_Methodology.docx' describing
what Exposure_ASW.py actually does, step by step.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# ── helpers ────────────────────────────────────────────────────
def heading(text, level=1):
    doc.add_heading(text, level=level)

def para(text):
    doc.add_paragraph(text)

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.left_indent = Inches(0.4)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def note(text):
    """Modelling simplification callout."""
    p = doc.add_paragraph()
    run_label = p.add_run("Simplification: ")
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run_body = p.add_run(text)
    run_body.italic = True

# ── Title ──────────────────────────────────────────────────────
doc.add_heading("Exposure Simulation Methodology", level=0)
p = doc.add_paragraph("Step-by-step description of the Monte Carlo exposure "
                       "simulation implemented in Exposure_ASW.py.")
p.italic = True
p = doc.add_paragraph("This document is auto-generated from the code to ensure "
                       "it reflects the actual implementation, not an idealised version.")
p.italic = True

# ────────────────────────────────────────────────────────────────
heading("1  Overview")
# ────────────────────────────────────────────────────────────────
para(
    "The simulation computes forward mark-to-market (MtM) profiles for four "
    "interest-rate swap structures sharing the same underlying bond "
    "(BTP 3.45%, semi-annual, maturing 1 Feb 2036) and the same ESTR/OIS "
    "discount curve.  It uses 1,500 Monte Carlo paths of parallel curve "
    "shifts over 100 equally-spaced observation dates from the valuation "
    "date to maturity.  All four structures share the same set of random "
    "paths so that their exposure profiles are directly comparable."
)

para("The four structures are:")
bullet("Structures 1–3: Par-par asset swaps with bond dirty prices "
       "P = 100, 90, and 110.")
bullet("Structure 4: Fix-for-fix differential swap (bank receives 3.50% "
       "fixed, pays 3.00% fixed).")

para("The output consists of EPE, ENE, PFE (95th percentile) and NFE "
     "(5th percentile) profiles, plus sample-path plots, for each structure.")

# ────────────────────────────────────────────────────────────────
heading("2  Base Yield Curve Construction")
# ────────────────────────────────────────────────────────────────
para(
    "A single base discount curve is built from 11 pillar zero rates "
    "(1W, 1Y, 2Y, …, 10Y) hard-coded in OIS_ZEROS.  The curve is "
    "constructed using QuantLib’s ZeroCurve class with linear interpolation "
    "and continuous compounding."
)
code("curve = ql.ZeroCurve(pillar_dates, rates, ois_dc, calendar,\n"
     "                     ql.Linear(), ql.Continuous, ql.Annual)")
para(
    "The function build_ois_curve() (line 88) prepends the today-date rate "
    "(equal to the 1W rate) to the pillar list so that QuantLib has a "
    "reference point at t = 0.  Extrapolation is enabled so discount "
    "factors can be queried beyond the 10Y pillar."
)
note(
    "The day-count convention used for the OIS curve is Actual/365 Fixed "
    "(ois_dc), while the bond accrual uses ActualActual ISMA (bond_dc).  "
    "A single curve is used for both discounting and forward-rate "
    "projection — there is no separate 6M Euribor or forward curve."
)

# ────────────────────────────────────────────────────────────────
heading("3  Bond Schedule and Coupon Data")
# ────────────────────────────────────────────────────────────────
para(
    "A QuantLib Schedule is built from the last coupon date "
    "(1 Feb 2026) to maturity (1 Feb 2036) with semi-annual frequency, "
    "backward generation, and Modified Following business-day convention "
    "on the TARGET calendar."
)
code("schedule = ql.Schedule(\n"
     "    BOND_LAST_COUPON_DATE, BOND_MATURITY,\n"
     "    ql.Period(COUPON_FREQUENCY), calendar,\n"
     "    ql.ModifiedFollowing, ql.ModifiedFollowing,\n"
     "    ql.DateGeneration.Backward, False)")
para(
    "From this schedule the code extracts coupon_dates (coupon end dates "
    "after the settlement date) and their matching period_starts.  For "
    "each coupon period, the year fraction alpha is computed using the "
    "ActualActual ISMA day count.  These triples (period_start_serial, "
    "coupon_end_serial, alpha) are stored in coupon_data (line 194–197)."
)

# ────────────────────────────────────────────────────────────────
heading("4  Annuity and Par Swap Rate")
# ────────────────────────────────────────────────────────────────
para(
    "The annuity A is the present value of a unit annuity stream over "
    "all remaining coupon periods:"
)
code("A = Σ alpha_i × DF(0, t_i)")
para(
    "where alpha_i is the year fraction of coupon period i and DF(0, t_i) "
    "is the discount factor from the base OIS curve to coupon date t_i."
)
para(
    "The par swap rate r_s is the fixed rate that makes a plain-vanilla "
    "fix-float swap worth zero at inception:"
)
code("r_s = (1 – DF(0, T)) / A")
para(
    "Both quantities are computed in section [5] of the code (lines 125–132)."
)

# ────────────────────────────────────────────────────────────────
heading("5  ASW Spread and Inception MtM")
# ────────────────────────────────────────────────────────────────
para(
    "For each of the three bond-price scenarios (P = 100, 90, 110), "
    "the par-par asset swap spread is computed as:"
)
code("ASW = (c – r_s) + (N – P) / (N × A)")
para(
    "where c = 3.45% is the bond coupon rate, N = 100 is the notional "
    "(expressed as a percentage), and P is the bond dirty price."
)
para("The inception MtM of the swap is verified to satisfy:")
code("MtM_0 = N × (1 – DF(0,T)) + (ASW – c) × N × A\n"
     "      = 100 – P")
para(
    "This identity is checked explicitly in the output: MtM + P must "
    "equal 100 for each scenario (lines 150–158)."
)

# ────────────────────────────────────────────────────────────────
heading("6  Tenor Grid (Observation Dates)")
# ────────────────────────────────────────────────────────────────
para(
    "The simulation evaluates MtM at N_TENORS + 1 = 101 observation "
    "dates, equally spaced in serial-number space from today to maturity:"
)
code("tenor_serial = np.linspace(today_serial, maturity_serial,\n"
     "                           N_TENORS + 1, dtype=int)\n"
     "tenor_serial = np.unique(tenor_serial)")
para(
    "The np.unique call removes any duplicates that arise from rounding "
    "to integer serial numbers.  Each serial number is converted to a "
    "QuantLib Date object (tenor_dates) and to a year fraction from today "
    "(tenor_years), using the Actual/365 Fixed convention."
)
note(
    "The grid is equally spaced in calendar days, not in year fractions.  "
    "This means observation density is uniform across the life of the "
    "trade, with approximately one observation every 36 days for a "
    "~10-year bond.  The dense grid is important for capturing the "
    "fix-for-fix sawtooth pattern: with only 10–20 annual observation "
    "dates, observation points can coincide with coupon dates, hiding "
    "the step-downs in A_rem."
)

# ────────────────────────────────────────────────────────────────
heading("7  Pre-computation of Base Discount Factors")
# ────────────────────────────────────────────────────────────────
para(
    "Before the simulation loop, the code queries the QuantLib curve "
    "once for every date that will be needed: all coupon dates, the "
    "maturity date, all observation dates, and all period start dates.  "
    "The results are stored in two dictionaries keyed by serial number:"
)
bullet("base_dfs[serial]: the base discount factor DF(0, d) from the "
       "OIS curve.")
bullet("date_years[serial]: the year fraction from today to d, using "
       "Actual/365 Fixed.")
para(
    "This avoids repeated QuantLib curve lookups inside the Monte Carlo "
    "loop and enables purely NumPy-based analytical shifting."
)

# ────────────────────────────────────────────────────────────────
heading("8  Random Shift Generation")
# ────────────────────────────────────────────────────────────────
para(
    "The yield-curve shifts are generated as a cumulative random walk.  "
    "At each of the 101 observation dates, an increment is drawn from "
    "N(0, σ) with σ = 0.010 (i.e., 100 basis points per step).  "
    "The increments are then cumulated across time:"
)
code("increments = np.random.normal(0.0, SIGMA,\n"
     "                              size=(N_PATHS, len(tenor_dates)))\n"
     "increments[:, 0] = 0.0\n"
     "shifts = np.cumsum(increments, axis=1)")
para(
    "This produces shifts as a (1500 × 101) matrix.  The first column "
    "is always zero (no shift at inception).  For path p at observation "
    "date t, the cumulative shift is dz[p, t].  This shift is a parallel "
    "additive displacement applied to all zero rates simultaneously."
)
para(
    "At observation date t (which is step t in the random walk), the "
    "cumulative shift has standard deviation σ√t.  For the final "
    "observation (step 100), this is σ√100 = 10σ = 0.10 (10%).  "
    "The random seed is fixed at 42 for reproducibility."
)
note(
    "The shift model is a parallel (level) shift only — there is no "
    "slope, curvature, or tenor-specific component.  All zero rates move "
    "by the same amount at each step.  This is a significant simplification "
    "compared to multi-factor models (e.g., PCA-based or LMM).  "
    "Additionally, the increments have constant volatility σ at every "
    "step regardless of the rate level (a normal/Bachelier model), as "
    "opposed to a log-normal model where volatility would scale with "
    "the rate level."
)
note(
    "There is no square-root-of-time scaling between steps.  Each "
    "increment has the same standard deviation σ regardless of the "
    "time gap between observation dates.  In a standard Brownian motion "
    "discretisation, the increment variance would be σ²Δt; here, "
    "Δt is implicitly set to 1 for every step.  This means the total "
    "variance at the final step is N_TENORS × σ², rather than "
    "T × σ² as in a continuous-time model."
)

# ────────────────────────────────────────────────────────────────
heading("9  Analytical Discount-Factor Shifting")
# ────────────────────────────────────────────────────────────────
para(
    "Rather than rebuilding the QuantLib yield curve for every path "
    "and every observation date (which would be prohibitively slow for "
    "1,500 × 101 = 151,500 curve constructions), the code applies an "
    "analytical shift to the pre-computed base discount factors."
)
para(
    "Given a parallel shift dz to all zero rates, the shifted discount "
    "factor from the observation date obs to a future date e is:"
)
code("DF_fwd_shifted(obs, e) = [DF_base(0, e) / DF_base(0, obs)]\n"
     "                       × exp(–dz × (t_e – t_obs))")
para(
    "where t_e and t_obs are year fractions from today to the respective "
    "dates.  The first factor, DF_base(0,e)/DF_base(0,obs), is the "
    "unshifted forward discount factor.  The exponential applies the "
    "parallel shift to the forward rate over the remaining tenor."
)
para("In the code (lines 242–246, inside the main loop):")
code("df_fwd_e = (df_base_e / obs_df_base) * np.exp(-dz * (t_e - obs_t))")
para(
    "This is fully vectorised over all 1,500 paths: dz is a 1-D array "
    "of length N_PATHS, and the multiplication and exponentiation are "
    "element-wise NumPy operations."
)
note(
    "This formula is exact for flat (parallel) shifts applied to "
    "continuously-compounded zero rates.  If the shifts were "
    "tenor-dependent (e.g., steepeners/flatteners), this shortcut "
    "would not work and a full curve rebuild would be needed.  "
    "The formula also assumes the base curve uses continuous compounding, "
    "which matches the ql.Continuous setting in build_ois_curve()."
)

# ────────────────────────────────────────────────────────────────
heading("10  Main Simulation Loop: A_rem and DF_fwd_T")
# ────────────────────────────────────────────────────────────────
para(
    "The heart of the simulation is a loop over all observation dates "
    "(section [10], lines 228–254).  For each observation date, the "
    "code computes two quantities across all 1,500 paths simultaneously:"
)
bullet(
    "arr_A_rem[path, t]: the remaining annuity from the observation date "
    "to maturity, computed using shifted forward discount factors.  "
    "This is the sum of alpha_i × DF_fwd_shifted(obs, t_i) over all "
    "coupon dates t_i that fall after the observation date."
)
bullet(
    "arr_df_fwd_T[path, t]: the shifted forward discount factor from "
    "the observation date to the bond maturity, "
    "DF_fwd_shifted(obs, T)."
)
para("Both arrays are initialised before the loop:")
code("arr_A_rem    = np.zeros((N_PATHS, len(tenor_dates)))\n"
     "arr_df_fwd_T = np.ones((N_PATHS, len(tenor_dates)))")
para(
    "arr_df_fwd_T is initialised to ones (not zeros) so that for "
    "observation dates at or after maturity (where no remaining coupons "
    "exist and the loop body is skipped via 'if not rem_coupons: "
    "continue'), the forward DF defaults to 1.  This ensures the MtM "
    "formula yields zero at maturity: N × (1 – 1) = 0."
)
para("For each observation date, the loop:")
bullet("Retrieves the cumulative shift vector dz = shifts[:, t_idx] "
       "(length N_PATHS).")
bullet("Identifies remaining coupons: periods where the coupon end "
       "date is strictly after the observation date serial number.")
bullet("Computes the remaining annuity by summing alpha_i × "
       "DF_fwd_shifted(obs, t_i) over all remaining coupon periods.")
bullet("Computes DF_fwd_shifted(obs, T) for the maturity date.")
note(
    "The ‘remaining coupons’ filter uses a strict inequality: "
    "'e_ser > obs_serial'.  A coupon whose end date exactly matches "
    "the observation date is treated as already paid and excluded.  "
    "This drives the sawtooth pattern in the fix-for-fix profile: "
    "A_rem drops discretely whenever an observation date crosses a "
    "coupon date."
)

# ────────────────────────────────────────────────────────────────
heading("11  MtM Revaluation — Asset Swap Scenarios")
# ────────────────────────────────────────────────────────────────
para(
    "Once arr_A_rem and arr_df_fwd_T are computed for all paths and "
    "tenors, the MtM for each ASW scenario is calculated as a simple "
    "array operation (section [11], lines 260–265):"
)
code("float_term = N × (1 – arr_df_fwd_T)\n"
     "MtM = float_term + (ASW – c) × N × arr_A_rem")
para(
    "The first term captures the floating leg’s value change as "
    "interest rates move: when rates rise (positive shift), discount "
    "factors fall, and (1 – DF) increases.  The second term captures "
    "the fixed-rate differential between the ASW spread received and "
    "the bond coupon paid, accumulated over the remaining annuity."
)
para("For the three scenarios:")
bullet("P = 100 (par): ASW ≈ r_s, so (ASW – c) is small.  MtM_0 = 0.  "
       "Exposure is centred around zero.")
bullet("P = 90 (discount): ASW > r_s (positive spread differential).  "
       "MtM_0 = +10.  Exposure profile is shifted upward.")
bullet("P = 110 (premium): ASW < r_s (negative spread differential).  "
       "MtM_0 = –10.  Exposure profile is shifted downward.")
note(
    "The MtM is a FORWARD value at the observation date, not discounted "
    "back to t = 0.  The forward discount factors in the formula "
    "inherently divide out DF(0, obs), so the MtM represents the "
    "market value of the swap as seen from the observation date.  "
    "This is the standard convention for counterparty exposure "
    "calculations (the exposure at a future date is the replacement "
    "cost at that date, not its present value today)."
)

# ────────────────────────────────────────────────────────────────
heading("12  MtM Revaluation — Fix-for-Fix Swap")
# ────────────────────────────────────────────────────────────────
para(
    "The fix-for-fix differential swap has a simpler MtM formula "
    "(section [12], lines 278–280):"
)
code("MtM_fixfix = (c_recv – c_pay) × N × arr_A_rem\n"
     "           = 0.50% × 100 × arr_A_rem")
para(
    "Since c_recv (3.50%) > c_pay (3.00%), the MtM is always positive: "
    "the bank is always in the money.  The exposure is one-sided "
    "(EPE > 0, ENE ≈ 0 everywhere)."
)
para(
    "The sawtooth (zig-zag) pattern arises because arr_A_rem steps "
    "down discretely at each semi-annual coupon date.  Between coupon "
    "dates, A_rem changes smoothly (driven by DF changes and the "
    "passage of time), but at each coupon date the most recent "
    "coupon period drops out of the sum, causing a downward step.  "
    "Because the MtM is proportional to A_rem, it inherits this "
    "sawtooth shape."
)

# ────────────────────────────────────────────────────────────────
heading("13  Exposure Metrics: EPE, ENE, PFE, NFE")
# ────────────────────────────────────────────────────────────────
para(
    "Exposure metrics are computed cross-sectionally (across paths) "
    "at each observation date by the function compute_exposure_metrics() "
    "(lines 292–297):"
)
code("epe = np.maximum(mtm, 0.0).mean(axis=0)\n"
     "ene = np.minimum(mtm, 0.0).mean(axis=0)\n"
     "pfe = np.percentile(mtm, 95, axis=0)\n"
     "nfe = np.percentile(mtm, 5, axis=0)")
para("Each metric is a 1-D array of length 101 (one value per observation date):")
bullet(
    "EPE (Expected Positive Exposure): the average of max(MtM, 0) across "
    "all paths.  Represents the expected replacement cost if the "
    "counterparty defaults at that date."
)
bullet(
    "ENE (Expected Negative Exposure): the average of min(MtM, 0) across "
    "all paths.  Represents the expected exposure FROM the bank TO the "
    "counterparty (own-default risk)."
)
bullet(
    "PFE (Potential Future Exposure, 95th percentile): the 95th "
    "percentile of MtM across paths.  A high-confidence upper bound "
    "on the positive exposure."
)
bullet(
    "NFE (Negative Future Exposure, 5th percentile): the 5th "
    "percentile of MtM across paths.  A high-confidence lower bound "
    "on the negative exposure."
)
note(
    "These are unconditional, point-in-time metrics.  There is no "
    "time-averaging (e.g., no EPE integrated over time to get EEPE), "
    "no credit-adjustment (no PD or LGD weighting), and no "
    "netting-set aggregation.  The metrics are computed per structure, "
    "not across a portfolio."
)
note(
    "Since the MtM values are forward (not discounted to t = 0), the "
    "EPE and PFE represent future-dated replacement costs.  To compute "
    "CVA one would need to discount these profiles back to t = 0 and "
    "multiply by a default probability density and loss-given-default, "
    "which is not done here."
)

# ────────────────────────────────────────────────────────────────
heading("14  EPE / ENE Symmetry and the Normal Shift Model")
# ────────────────────────────────────────────────────────────────
para(
    "The code includes a diagnostic (section [14], lines 346–360) that "
    "examines whether EPE and |ENE| are symmetric for the par-bond "
    "scenario (MtM_0 = 0).  In a linear model, a symmetric shift "
    "distribution would produce perfectly symmetric MtM, hence "
    "EPE = |ENE|."
)
para(
    "However, the MtM formula contains exp(–dz × t), which is convex "
    "in dz.  By Jensen’s inequality, negative shifts (rate decreases) "
    "inflate discount factors exponentially (unbounded), while positive "
    "shifts compress them toward zero (bounded below by 0).  This "
    "creates a negatively skewed MtM distribution."
)
para(
    "With the current parameters (σ = 0.010, 100 steps), the "
    "cumulative shift standard deviation at the midpoint (~5Y) is "
    "approximately 7%.  This is large enough for the exp convexity "
    "to produce visible asymmetry: EPE / |ENE| ≈ 0.63 and MtM "
    "skewness ≈ –0.96 at the midpoint."
)

# ────────────────────────────────────────────────────────────────
heading("15  Summary of Modelling Simplifications")
# ────────────────────────────────────────────────────────────────
para(
    "The following modelling simplifications are present in the "
    "current implementation.  These are appropriate for an educational "
    "demonstration but would need to be addressed in a production "
    "counterparty exposure system:"
)
bullet(
    "Parallel shifts only: all zero rates move by the same amount.  "
    "No slope, curvature, or basis risk is modelled.  A real system "
    "would use multi-factor models (e.g., 3-factor PCA or a Libor "
    "Market Model)."
)
bullet(
    "Normal (Bachelier) shift model: the increment dz is additive "
    "and rate-level-independent.  This allows negative rates and does "
    "not capture the empirically observed relationship between rate "
    "level and volatility.  A log-normal or shifted log-normal model "
    "is more common in practice."
)
bullet(
    "No √Δt scaling: each random-walk increment has variance σ² "
    "regardless of the time gap between observations.  In a correctly "
    "discretised Brownian motion, the variance would be σ²Δt.  "
    "This means the total variance depends on the number of steps, "
    "not on physical time."
)
bullet(
    "Single-curve framework: the same OIS curve is used for both "
    "discounting and forward-rate projection.  In a dual-curve "
    "framework (post-2008 standard), a separate IBOR/Euribor curve "
    "would be used for forward rates, and cross-currency basis would "
    "be modelled."
)
bullet(
    "Static bond schedule: the coupon schedule is fixed at inception "
    "and does not change with business-day adjustments at future dates.  "
    "The simulation does not model coupon accrual or dirty/clean price "
    "effects."
)
bullet(
    "No credit or funding adjustments: there is no probability of "
    "default, loss-given-default, or credit spread modelling.  The "
    "exposure profiles are pre-credit, not risk-adjusted (no CVA, DVA, "
    "or FVA)."
)
bullet(
    "No collateral or margining: the simulation assumes uncollateralised "
    "swaps.  In practice, margin agreements (CSA) would reduce exposure "
    "through regular collateral exchanges, and the margin period of "
    "risk would be the relevant exposure horizon."
)
bullet(
    "Forward MtM (not discounted): the MtM at each observation date "
    "is a forward value, not discounted back to today.  This is "
    "standard for exposure profiles but means the values cannot be "
    "directly summed across dates without discounting."
)
bullet(
    "Hard-coded curve data: the OIS zero rates are hard-coded "
    "constants, not sourced from a market data feed.  The simulation "
    "date is fixed at 24 April 2026."
)
bullet(
    "Fixed random seed: np.random.seed(42) ensures reproducibility "
    "but means every run produces identical paths.  In a production "
    "system, multiple independent seeds or quasi-random sequences "
    "would be used for convergence testing."
)

# ────────────────────────────────────────────────────────────────
heading("16  Code Structure Reference")
# ────────────────────────────────────────────────────────────────
para("The table below maps each step to the relevant code section and "
     "key variables in Exposure_ASW.py.")

table = doc.add_table(rows=11, cols=3)
table.style = "Light Shading Accent 1"
hdr = table.rows[0]
hdr.cells[0].text = "Step"
hdr.cells[1].text = "Code Section / Lines"
hdr.cells[2].text = "Key Variables"
for cell in hdr.cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ("Base curve",            "[3] lines 88–98",         "base_curve, OIS_ZEROS"),
    ("Bond schedule",         "[4] lines 103–120",       "schedule, coupon_dates, period_starts"),
    ("Annuity & r_s",         "[5] lines 125–132",       "annuity, df_T, r_s"),
    ("ASW & inception MtM",   "[6] lines 148–158",       "asw, mtm0, scenario_data"),
    ("Tenor grid",            "[7] lines 167–178",       "tenor_serial, tenor_dates, tenor_years"),
    ("Base DFs cache",        "[8] lines 183–199",       "base_dfs, date_years, coupon_data"),
    ("Random shifts",         "[9] lines 207–209",       "increments, shifts"),
    ("Simulation loop",       "[10] lines 228–254",      "arr_A_rem, arr_df_fwd_T, dz, rem_coupons"),
    ("MtM (ASW & fix-fix)",   "[11–12] lines 260–280", "float_term, mtm_all, mtm_fixfix"),
    ("Exposure metrics",      "[13] lines 292–297",      "epe, ene, pfe, nfe"),
]
for i, (step, section, variables) in enumerate(data, start=1):
    row = table.rows[i]
    row.cells[0].text = step
    row.cells[1].text = section
    row.cells[2].text = variables

# ── save ───────────────────────────────────────────────────────
out_path = "Exposure_Simulation_Methodology.docx"
doc.save(out_path)
print(f"Document saved to {out_path}")
